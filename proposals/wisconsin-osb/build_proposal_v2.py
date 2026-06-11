#!/usr/bin/env python3
"""Build the v2 (alternate) draft of the Victor-Strategies Wisconsin OSB proposal.

A separate draft that restructures the REVISED 061226 document for persuasion:
executive summary led by independence, qualifications mapped to phases with a
Choctaw case study, forecast track record vs. actual market outcomes, an
illustrative GEMS table for Wisconsin, and a fact-based summary. The full
chronological engagement history moves to an appendix. Unchanged sections are
reused from build_proposal.py.

Regenerate with:
    python3 build_proposal_v2.py     (requires `pip install python-docx`)

Output: VS_Proposal_Online_Sports_Betting_Wisconsin_Data_Analysis_REVISED_v2_061226.docx
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import build_proposal as bp

OUTPUT = os.path.join(
    bp.HERE,
    "VS_Proposal_Online_Sports_Betting_Wisconsin_Data_Analysis_REVISED_v2_061226.docx",
)

WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------- v2 helpers

def callout_box(doc, title, body_paragraphs):
    """Single-cell shaded callout box with a red title line."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bp.light_table_borders(table)
    bp.fixed_layout(table, (Inches(6.5),))
    cell = table.rows[0].cells[0]
    bp.shade_cell(cell, "F7F2F2")
    p = cell.paragraphs[0]
    bp.set_font(p.add_run(title), font=bp.HEAD_FONT, size=11, bold=True,
                color=bp.VS_RED)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for text_or_runs in body_paragraphs:
        body = cell.add_paragraph()
        if isinstance(text_or_runs, str):
            text_or_runs = [(text_or_runs, False, False)]
        for text, bold, italic in text_or_runs:
            bp.set_font(body.add_run(text), size=10, bold=bold, italic=italic)
        body.paragraph_format.space_after = Pt(6)
        body.paragraph_format.line_spacing = 1.15
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def data_table(doc, header, rows, col_widths, align_numbers=True):
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bp.light_table_borders(table)
    bp.repeat_header_row(table)
    hdr = table.rows[0].cells
    for i, label in enumerate(header):
        bp.shade_cell(hdr[i], "C42127")
        p = hdr[i].paragraphs[0]
        bp.set_font(p.add_run(label), font=bp.HEAD_FONT, size=10, bold=True,
                    color=WHITE)
        p.paragraph_format.space_after = Pt(2)
        if align_numbers and i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row_values in rows:
        row = table.add_row()
        bp.keep_row_together(row)
        for i, (cell, text) in enumerate(zip(row.cells, row_values)):
            p = cell.paragraphs[0]
            bp.set_font(p.add_run(text), size=9.5, bold=(i == 0))
            p.paragraph_format.space_after = Pt(2)
            if align_numbers and i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    bp.fixed_layout(table, col_widths)
    return table


def cover_page_v2(doc):
    for _ in range(3):
        doc.add_paragraph()
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.add_run().add_picture(os.path.join(bp.ASSETS, "vs-logo.png"),
                               width=Inches(3.6))
    for _ in range(2):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.set_font(title.add_run("Online Sports Betting in the State of Wisconsin"),
                font=bp.HEAD_FONT, size=26, bold=True, color=bp.VS_DARK)
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.set_font(sub.add_run("Data Analysis"), font=bp.HEAD_FONT, size=18,
                bold=True, color=bp.VS_RED)
    sub.paragraph_format.space_after = Pt(2)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.set_font(rule.add_run("—  ❖  —"), font=bp.HEAD_FONT, size=11,
                color=bp.VS_GRAY)

    doc.add_paragraph()
    prep = doc.add_paragraph()
    prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.set_font(prep.add_run("Prepared for"), font=bp.HEAD_FONT, size=12,
                italic=True, color=bp.VS_GRAY)
    prep.paragraph_format.space_after = Pt(2)

    client = doc.add_paragraph()
    client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.set_font(client.add_run("Wisconsin Tribal Nations\nAdvisory Working Group"),
                font=bp.HEAD_FONT, size=16, bold=True, color=bp.VS_DARK)

    doc.add_paragraph()
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.set_font(date.add_run("June 12, 2026"), font=bp.HEAD_FONT, size=12,
                color=bp.VS_DARK)

    mark = doc.add_paragraph()
    mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.set_font(mark.add_run("DRAFT 2 — CONFIDENTIAL"), font=bp.HEAD_FONT,
                size=10, bold=True, color=bp.VS_RED, caps=True)

    doc.add_page_break()


# ---------------------------------------------------------------- new sections

def executive_summary(doc):
    bp.heading1(doc, "Executive Summary")
    bp.para(doc,
            "Victor-Strategies proposes to serve as the Working Group's "
            "independent analyst for the online sports betting models under "
            "review: validating the data and market assumptions behind each "
            "model (Phase I), comparing the revenue-sharing structures on a "
            "consistent, per-tribe basis over ten years (Phase II), and "
            "assessing the prediction-market risk to tribal exclusivity "
            "(Phase III). All three phases are completed within 60 days of "
            "contract signing.")

    callout_box(doc, "Independence", [
        "We have no current relationships with any of the 11 Wisconsin tribes "
        "and no current financial relationship with proponents of either model "
        "under review. Our only stake in this engagement is the accuracy of "
        "the analysis. The Working Group's referee should have nothing to "
        "gain from the answer — and we don't.",
    ])

    bp.para(doc, [("Why Victor-Strategies:", True, False)], space_after=4)
    bp.bullet_lead(doc, "We have done this engagement before. ",
                   "Over three years (2018–2021) we advised the Choctaw Nation "
                   "of Oklahoma on its sports betting initiatives, delivering "
                   "the same components this RFP requests: a regional market "
                   "study, a 10-year online sports betting revenue model, a "
                   "retail forecast, and an economic impact analysis (case "
                   "study, Section A).")
    bp.bullet_lead(doc, "Our public forecasts have been tested by real markets. ",
                   "Our team's 2018 statewide study for the Indiana Gaming "
                   "Commission projected a $256 million annual market at "
                   "maturity and recommended the mobile-led, low-tax structure "
                   "Indiana adopted in 2019. The market has since confirmed "
                   "the forecast's structure — and its deliberately "
                   "conservative posture (Section A).")
    bp.bullet_lead(doc, "We know the bettor. ",
                   "Since the repeal of PASPA in May 2018, Victor-Strategies "
                   "has interviewed more than 25,000 sports bettors across its "
                   "client research projects.")
    bp.bullet_lead(doc, "We cover the risk this RFP is most concerned about. ",
                   "Dustin Gouker is arguably the industry's most informed "
                   "analyst of prediction markets, and our team has monitored "
                   "the U.S. wagering market monthly since 2023 for a major "
                   "international platform provider, with particular attention "
                   "to tribal government gaming.")
    bp.bullet_lead(doc, "Purpose-built economic modeling. ",
                   "The optional statewide economic impact study runs on GEMS, "
                   "a gaming-specific input-output platform with dedicated "
                   "online-gambling multipliers that off-the-shelf software "
                   "lacks, backed by more than 165 engagements since 2005 "
                   "(Sections D and J).")
    bp.bullet_lead(doc, "Our work holds up under attack. ",
                   "When a funded opposing study disputed our New Jersey "
                   "iGaming findings, we tested both models line by line and "
                   "published the analysis. Arbitrating between competing "
                   "economic models — under scrutiny — is precisely the role "
                   "this engagement requires.")

    doc.add_heading("Engagement Timeline", level=2)
    data_table(
        doc,
        ("Phase", "Scope", "Completion"),
        [
            ("Phase I", "Market assessment and data validation",
             "Day 45"),
            ("Phase II", "Revenue share model review and comparison",
             "Day 45 (parallel with Phase I)"),
            ("Phase II option", "Statewide economic impact study "
             "(priced separately)", "Aligned with Phase II"),
            ("Phase III", "Prediction markets", "Day 60"),
        ],
        (Inches(1.2), Inches(3.4), Inches(2.2)),
        align_numbers=False,
    )
    bp.para(doc,
            "Total elapsed time: 60 days from contract signing, with Phases I "
            "and II running in parallel.",
            space_after=0)


def section_a_v2(doc):
    bp.heading1(doc, "A. Company Overview and Relevance of Work")

    doc.add_heading("Company Overview", level=2)
    bp.para(doc,
            "Victor-Strategies is a professional advisory firm providing expert "
            "services and critical business insights to the gaming industry, with a "
            "particular focus on Indian Country. We provide governments and industry "
            "leaders with the strategies, tools, data, and expertise essential for "
            "informed decision making and effective operations, supported by "
            "research and analysis capabilities of the highest quality.")
    bp.para(doc,
            "Our business model is built on trusted-advisor relationships through "
            "which clients can draw on our broad network of strategic partners. That "
            "network offers expertise in tribal and commercial gaming — both "
            "land-based and interactive — as well as government policy and media "
            "strategy, and economic impact and development through conventional and "
            "emerging technologies.")
    bp.para(doc,
            "Our principals are recognized thought leaders in the gaming industry "
            "and are committed to delivering maximum value to our clients. Our "
            "founder, Victor Rocha, serves as Conference Chairman of the Indian "
            "Gaming Association while leading Victor-Strategies as its president. He "
            "is also the owner and publisher of Pechanga.net, an online information "
            "resource for Native American tribes across North America, and has been "
            "deeply engaged in the political landscape of U.S. tribal gaming since "
            "1998. Gene Johnson brings more than 35 years of gaming-industry "
            "experience spanning information technology, business process "
            "improvement, quality assurance, and strategic planning and analysis.")
    bp.para(doc,
            "Since its founding in 2016, Victor-Strategies has conducted primary and "
            "secondary research for scores of clients in international markets and "
            "in more than thirty U.S. states. This research includes market sizing, "
            "revenue analysis, feasibility studies, economic impact, player "
            "behavior, gaming consumer research, and new gaming technology. Since "
            "the overturn of the Professional and Amateur Sports Protection Act "
            "(PASPA) in May 2018, Victor-Strategies has cumulatively interviewed "
            "more than 25,000 sports bettors across its client research projects. "
            "Selected publicly available examples of our work include the "
            "following:")
    for url in (
        "https://ideagrowth.org/research/",
        "https://ideagrowth.org/wp-content/uploads/2020/01/Economic-Impact-of-NJ-iGaming_FULL-REPORT_12.19.19.pdf",
        "https://www.in.gov/igc/files/Indiana-SportsBettingReport-Final-Oct18-1.pdf",
        "https://sbcamericas.com/tag/the-more-you-know-the-bettor-defining-us-sports-betting-customer-personas/",
        "http://www.indiangaming.org/info/alerts/Spectrum-Internet-Paper.pdf",
    ):
        bp.link_bullet(doc, url)

    doc.add_heading("Independence and Conflicts", level=2)
    bp.para(doc,
            "We have no current relationships with any of the 11 Wisconsin "
            "tribes and no current financial relationship with proponents of "
            "either model under review. We are familiar with IGRA Class III "
            "compact structures and tribal revenue-sharing frameworks, and most "
            "of our work is confidential and performed on behalf of tribal and "
            "government clients. For a validation engagement, we regard "
            "independence as the threshold qualification: our conclusions are "
            "worth only as much as our distance from the outcome.")

    doc.add_heading("Experience Mapped to This Engagement", level=2)
    bp.para(doc,
            "Rather than recite our engagement history chronologically, we "
            "summarize below the experience most relevant to each phase of the "
            "proposed scope. A complete chronological engagement history "
            "appears in Appendix K.")

    doc.add_heading("Phase I — Market Sizing and Forecast Validation", level=3)
    bp.bullet_lead(doc, "Minnesota (since 2022, and 2026). ",
                   "Ongoing advisory to a tribal gaming client on the "
                   "legislative outlook, revenue opportunities, and operational "
                   "requirements for legalized sports wagering — including "
                   "multiple market-sizing, handle, and revenue estimates for "
                   "mobile and retail betting, adjusted to each session's "
                   "legislative proposals — plus a 2026 feasibility study with "
                   "statewide and property-specific five-year forecasts. "
                   "Minnesota is the closest demographic and structural "
                   "comparable to Wisconsin in our portfolio.")
    bp.bullet_lead(doc, "New York (2017 and 2025) and Arkansas (2019). ",
                   "Statewide and property-specific handle and revenue "
                   "forecasts, including a 10-year forecast for one of the "
                   "successful bidders in the New York City casino license "
                   "competition.")
    bp.bullet_lead(doc, "Proprietary market data. ",
                   "Alfonso Straffon maintains a proprietary database of "
                   "online sports betting statistics used by Wall Street "
                   "investment firms; since 2018 we have also partnered with "
                   "Eilers & Krejcik Gaming on national and statewide market "
                   "sizing and consumer research.")
    p = bp.para(doc,
                "Where our public forecasts can be tested against outcomes, "
                "they have held up. Our team's 2018 statewide study for the "
                "Indiana Gaming Commission, prepared with Eilers & Krejcik "
                "Gaming, projected a $256 million annual market at five-year "
                "maturity and recommended a mobile-led market structure at a "
                "single-digit tax rate. Indiana adopted that framework in "
                "2019; the mobile channel now produces the large majority of "
                "handle, as forecast, and the market reached $404 million in "
                "gross revenue in calendar 2023 — ahead of our maturity "
                "estimate. We build forecasts to be conservative. That is the "
                "posture a validation engagement requires, and the Working "
                "Group should expect it from us here.")
    bp.superscript(p, "1")

    doc.add_heading("Phase II — Revenue-Sharing Structures and Model "
                    "Arbitration", level=3)
    callout_box(doc, "Case Study — Choctaw Nation of Oklahoma (2018–2021)", [
        "Over three years, Victor-Strategies advised the Choctaw Nation of "
        "Oklahoma on its white-label sports betting initiatives, delivering "
        "the same components this RFP requests: a regional market study, an "
        "online sports betting model with a 10-year revenue forecast, a "
        "retail sports betting forecast, and an economic impact analysis. We "
        "also surveyed more than 3,000 database customers and non-customers "
        "across Oklahoma and Texas, profiling betting behavior, attitudes "
        "toward legalization, and demographics — and demonstrating that a "
        "significant share of casino players were already betting on sports "
        "before legalization.",
        "This engagement is the closest analogue to the Working Group's "
        "requirement: a tribal government weighing sports betting market "
        "entry, with multi-year revenue forecasts, alternative operating "
        "structures, and economic impacts quantified side by side.",
    ])
    bp.bullet_lead(doc, "Arbitrating between competing models. ",
                   "In 2024, with Alan Meister of Meister Economic Consulting, "
                   "Doug Walker, and Dan Waugh of Regulus Partners, we "
                   "produced A Comprehensive Analysis of NERA's Study on New "
                   "Jersey's iGaming Economic Impact — testing two competing "
                   "economic models' data, assumptions, and methods line by "
                   "line and publishing the findings under adversarial "
                   "scrutiny. That referee role — evaluating others' models "
                   "and defending the conclusions — is precisely what Phase II "
                   "asks of the selected consultant.")
    bp.bullet_lead(doc, "Balancing operator and non-operator interests. ",
                   "As lead consultant to the Massachusetts State Lottery "
                   "Commission's Online Products Task Force (2012), Mr. "
                   "Johnson advised on adding an online channel while "
                   "protecting 7,400 retail outlets and the value of pending "
                   "casino licenses — structurally analogous to weighing "
                   "operating-tribe risk against non-operating-tribe "
                   "distributions in Wisconsin.")
    bp.bullet_lead(doc, "Tribal consortium policy framing. ",
                   "In 2013, our team analyzed the California Internet poker "
                   "market for a major tribal casino consortium — market "
                   "sizing, projected revenues by vertical, and optimal "
                   "taxation rates — used to frame policy decisions by a "
                   "consortium of tribal governments.")

    doc.add_heading("Phase III — Prediction Markets and Emerging Verticals",
                    level=3)
    bp.bullet_lead(doc, "Dedicated prediction-market expertise. ",
                   "Dustin Gouker, who helped build Legal Sports Report into "
                   "the industry's reference source, now publishes The Event "
                   "Horizon, a newsletter tracking prediction markets and "
                   "their collision with regulated sports betting.")
    bp.bullet_lead(doc, "Continuous market monitoring. ",
                   "From 2023 to 2025 we produced monthly reports on the U.S. "
                   "sports wagering market — including prediction-market "
                   "developments — for a major international provider of "
                   "sports betting data, platforms, and player account "
                   "management systems, with particular attention to tribal "
                   "government gaming.")
    bp.bullet_lead(doc, "Emerging-vertical advisory. ",
                   "From 2018 to 2020 we advised a domestic micro-betting "
                   "start-up on its U.S. rollout strategy, an engagement that "
                   "concluded when the client was acquired by a major U.S. "
                   "sports betting operator.")

    notes = doc.add_paragraph()
    bp.set_font(notes.add_run("Notes"), font=bp.HEAD_FONT, size=10, bold=True,
                color=bp.VS_GRAY)
    notes.paragraph_format.space_before = Pt(14)
    notes.paragraph_format.keep_with_next = True
    bp.bottom_rule(notes, color="D9D9D9", size="4")
    n1 = doc.add_paragraph()
    n1.paragraph_format.keep_together = True
    bp.superscript(n1, "1 ")
    bp.set_font(n1.add_run(
        "Forecast: Eilers & Krejcik Gaming (2018), Regulated Sports Betting in "
        "Indiana, prepared for the Indiana Gaming Commission, "), size=9)
    bp.add_hyperlink(n1, "https://www.in.gov/igc/files/Indiana-SportsBettingReport-Final-Oct18-1.pdf", size=9)
    bp.set_font(n1.add_run(
        ". Actuals: Indiana Gaming Commission monthly revenue reports, "), size=9)
    bp.add_hyperlink(n1, "https://www.in.gov/igc/publications/monthly-revenue/", size=9)
    bp.set_font(n1.add_run(
        " (calendar-2023 taxable adjusted gross revenue of $404.4 million)."),
        size=9)


def section_d_v2(doc):
    bp.heading1(doc, "D. Phase II Option — Statewide Economic Impact Study",
                new_page=True)

    doc.add_heading("Background", level=2)
    bp.para(doc,
            "Economic impact analysis was not specifically requested in the "
            "Wisconsin Tribal Nations request for proposal. During the question "
            "period, however, Victor-Strategies asked whether statewide economic "
            "impact was an element of any of the models under review, and whether "
            "the consultant should include a statewide economic impact component in "
            "its independent revenue model. The answer was affirmative, so we "
            "include an economic impact component in this response. It is offered as "
            "an additional option, priced separately from the other phases of our "
            "proposal.")
    p = bp.para(doc,
                "Before the methodology, one finding deserves to be stated "
                "plainly, because it answers the question tribal "
                "decision-makers ask first: will online sports betting take "
                "revenue from tribal casino floors? The peer-reviewed "
                "literature — including research by our lead economist — "
                "finds that online and land-based gambling behave as "
                "complements when the offerings are horizontally integrated: "
                "online wagering is associated with gains, not losses, in "
                "land-based gaming revenue. We treat this as an empirical "
                "question to be verified with Wisconsin data rather than an "
                "assumption, but the evidence runs in the tribes' favor.")
    bp.superscript(p, "2")

    doc.add_heading("Approach and Methodology", level=2)
    bp.para(doc,
            "If included, this macroeconomic analysis will estimate the total "
            "statewide economic impact of online sports wagering as forecast by the "
            "recommended sports betting revenue model. The impact of this new "
            "economic vertical will be measured across several dimensions: economic "
            "output; value added at the state level, or gross state product (GSP); "
            "full-time-equivalent (FTE) employment; employee compensation; and "
            "non-gaming tax revenue. Results will be reported as direct, indirect, "
            "and induced impacts to the State of Wisconsin and, by inference, to the "
            "tribes.")
    bp.para(doc,
            "Economic impact is a measure of the spending and employment associated "
            "with a business, a sector of the economy, a specific project (such as "
            "the construction of a new facility), or a change in government policy "
            "or regulation. It can be measured in several ways, most commonly the "
            "dollar value of output produced and the person-years (full-time "
            "equivalents) of employment generated. These figures assess the gross "
            "level of activity or expenditure; they are not net measures that weigh "
            "benefits against costs, but they are useful in developing an "
            "appreciation of businesses, projects, investments, and economic "
            "sectors. In our modeling we do account for the shifts in consumption "
            "that are clearest to define — namely, shifts by Wisconsin residents "
            "from other household goods and services to sports wagering.")

    doc.add_heading("The GEMS Modeling Platform", level=3)
    bp.para(doc,
            "Our calculations will be performed with GEMS (Gaming Economic Modeling "
            "System), a proprietary input-output (IO) modeling platform developed by "
            "GP Consulting and purpose-built for the gaming industry. GEMS covers "
            "all 50 states and is constructed from federal economic accounts: state "
            "IO multipliers derived from U.S. EPA state input-output models, "
            "employment and wage coefficients from the Bureau of Labor Statistics "
            "Quarterly Census of Employment and Wages, and industry detail from the "
            "Bureau of Economic Analysis (BEA).")
    bp.para(doc,
            "GEMS addresses a structural weakness of standard, off-the-shelf "
            "economic impact software for gaming applications. Conventional "
            "platforms model casinos within a blended industry sector that combines "
            "gambling with golf courses, fitness centers, and amusement parks — "
            "activities with fundamentally different labor and cost structures — "
            "which systematically misstates the employment and value-added effects "
            "of gaming operations. GEMS isolates gambling-specific multipliers from "
            "the BEA detail accounts and adds dedicated online-gambling coefficients "
            "that conventional models lack entirely. For an online sports betting "
            "analysis, this produces materially more accurate estimates than "
            "generic, blended-sector multipliers.")
    bp.para(doc,
            "GEMS produces transparent, replicable estimates of output, value added, "
            "employment, wages, and state and local fiscal effects, decomposed into "
            "direct, indirect, and induced impacts using Type I and Type II "
            "(Leontief) multipliers. It incorporates direct employment data where "
            "available and enables rapid, consistent comparison of alternative "
            "revenue scenarios — a useful feature when evaluating the alternative "
            "market structures contemplated in this engagement.")

    doc.add_heading("Illustrative GEMS Output for Wisconsin", level=3)
    bp.para(doc,
            "To make the option concrete, the table below shows what the "
            "platform produces today for the Wisconsin online gambling "
            "sector, scaled to each $100 million of gross gaming revenue "
            "retained in-state (2023 multiplier vintage, Type II):")
    data_table(
        doc,
        ("Measure (per $100M GGR)", "Direct", "Indirect", "Induced", "Total"),
        [
            ("Economic output ($M)", "100.0", "25.7", "49.8", "175.5"),
            ("Value added / GSP ($M)", "47.3", "12.1", "29.5", "88.9"),
            ("Employment (FTE jobs)", "153", "27", "60", "240"),
            ("Employee compensation ($M)", "12.6", "4.8", "11.3", "28.8"),
            ("State & local production taxes ($M)", "5.2", "0.7", "2.4", "8.3"),
        ],
        (Inches(2.9), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.9)),
    )
    bp.para(doc, [
        ("These figures are illustrative of the platform's current output, "
         "not a forecast. ", False, True),
        ("Engagement results will be driven by the Phase I revenue model and "
         "Wisconsin-specific operating assumptions — in particular the share "
         "of platform employment located in-state, promotional spending "
         "treatment, and the household-spending substitution assumptions "
         "described below.", False, True),
    ])

    doc.add_heading("Measures of Impact", level=3)
    bp.para(doc,
            "The statewide economic impact of sports wagering expansion will be "
            "measured across the following dimensions:")
    bp.bullet(doc, "Economic output")
    bp.bullet(doc, "Value added, or gross state product (GSP)")
    bp.bullet(doc, "Full-time-equivalent (FTE) employment")
    bp.bullet(doc, "Employee compensation")
    bp.bullet(doc, "Non-gaming tax revenue (production, payroll, and household "
                   "taxes)")
    bp.para(doc,
            "We estimate economic activity through three layers of effects — "
            "direct, indirect, and induced:")
    bp.bullet_lead(doc, "Direct economic impact ",
                   "is the employment and economic output attributable to the "
                   "operation and management of the sportsbook within the operator's "
                   "business. This includes the jobs located directly at the "
                   "operator, together with the professional staff, management, "
                   "marketing, and other workers required to run the platform.")
    bp.bullet_lead(doc, "Indirect economic impact ",
                   "is the employment, value added, and economic output created in "
                   "industries that supply goods and services to the operator — for "
                   "example, a local IT company that installs hardware, or an "
                   "in-state office supplier that furnishes back-of-house spaces.")
    bp.bullet_lead(doc, "Induced economic impact ",
                   "is the employment, value added, and economic output generated by "
                   "the spending of individuals employed directly or indirectly by "
                   "the operator — for example, operator employees consuming "
                   "services such as a meal at a Wisconsin restaurant.")
    bp.bullet_lead(doc, "Total economic impact ",
                   "is the sum of direct, indirect, and induced effects. The "
                   "multiplier (indirect and induced) impacts represent the maximum "
                   "potential stimulus to the economy from gaming-related business "
                   "activity.")

    doc.add_heading("Catalytic Impacts", level=3)
    p = bp.para(doc,
                "Total economic impacts are conventionally defined as the sum of "
                "direct, indirect, and induced effects. Some sectors, however, are "
                "also capable of producing catalytic impacts — economic growth "
                "enabled by another sector, industry, or firm. Aviation is often "
                "cited as a source of catalytic impacts because it allows trade that "
                "would otherwise not occur, such as air cargo transport of "
                "perishable goods.")
    bp.superscript(p, "3")
    bp.para(doc,
            "In Wisconsin, the expansion of sports betting could similarly "
            "facilitate catalytic impacts. If sports betting attracts new "
            "entrepreneurs to the state, for example, the industry may spur the "
            "development of a new set of businesses that would not otherwise have "
            "considered expanding there.")
    bp.para(doc,
            "Under a typical, conservative economic impact methodology, such "
            "potential catalytic impacts are not measured. We considered including "
            "catalytic estimates, but they would require significant assumptions "
            "about innovation and consumer behavior that would materially increase "
            "the margin of error on our projections. We are nonetheless comfortable "
            "asserting that adopting a more business-friendly set of policies will "
            "typically have a net positive effect on surrounding businesses, and we "
            "note that our figures and approach are conservative: there is reason to "
            "believe the Wisconsin economy would experience impacts beyond those "
            "specifically measured in this study.")

    doc.add_heading("Evidence of Impacts on Existing Gaming Businesses", level=3)
    bp.para(doc,
            "A common public policy question is whether new gaming revenue will "
            "cannibalize existing businesses in the area — in particular, other "
            "gaming businesses. The question is especially important when the "
            "potentially cannibalized business produces higher margins or public "
            "revenue. As noted under Background, the academic literature has "
            "consistently pointed to a non-negative effect: offline gaming "
            "businesses are estimated to realize positive revenue gains from "
            "increased sports betting participation when the systems can be "
            "horizontally integrated.")
    bp.para(doc,
            "An important test of whether economic impact figures can be trusted "
            "for policymaking is whether a negative impact is occurring elsewhere "
            "in the economy. While we expect aggregate effects on existing gaming "
            "businesses to be positive or nil (not cannibalistic), we make "
            "explicit assumptions about changes in household discretionary "
            "spending by Wisconsin residents: we assume that revenue from "
            "Wisconsin residents is diverted from other spending, based on "
            "existing spending patterns and scenario-based analyses.")

    doc.add_heading("Economic Impact Deliverables", level=2)
    bp.bullet(doc,
              "A separate section of our management summary report detailing the "
              "statewide direct, indirect, induced, and total economic impacts of "
              "an online sports betting industry across the tribal and commercial "
              "economies of Wisconsin.")
    bp.timeline(doc, "Priced separately as an optional addition to Phase II; "
                     "schedule to be aligned with the Phase II deliverables.")

    notes = doc.add_paragraph()
    bp.set_font(notes.add_run("Notes"), font=bp.HEAD_FONT, size=10, bold=True,
                color=bp.VS_GRAY)
    notes.paragraph_format.space_before = Pt(14)
    notes.paragraph_format.keep_with_next = True
    bp.bottom_rule(notes, color="D9D9D9", size="4")
    n2 = doc.add_paragraph()
    n2.paragraph_format.keep_together = True
    n2.paragraph_format.keep_with_next = True
    bp.superscript(n2, "2 ")
    bp.set_font(n2.add_run(
        "Philander, K., Abarbanel, B., & Repetti, T. (2015). Consumer spending "
        "in the gaming industry: Evidence of complementary demand in casino and "
        "online venues. International Gambling Studies. "), size=9)
    bp.add_hyperlink(n2, "https://www.tandfonline.com/doi/abs/10.1080/14459795.2015.1042002", size=9)
    n2.paragraph_format.space_after = Pt(3)
    n3 = doc.add_paragraph()
    n3.paragraph_format.keep_together = True
    bp.superscript(n3, "3 ")
    bp.set_font(n3.add_run(
        "Tretheway, M. (2010). Economic Impacts of Aviation: Catalytic Impacts. "
        "ACED Conference. "), size=9)
    bp.add_hyperlink(n3, "http://www.intervistas.com/downloads/presentations/Economic_Impacts_of_Aviation_Catalytic_Impacts_MTR_20Sep2010.pdf", size=9)


def section_f_v2(doc):
    bp.heading1(doc, "F. Financial Proposal (RFP §8.4)")
    bp.para(doc,
            "Per the RFP, respondents must provide a detailed fee schedule "
            "including:")
    bp.bullet(doc, "Estimated hours by staff level, with hourly rates and total "
                   "cost.")
    bp.bullet(doc, "Out-of-pocket expenses, including a not-to-exceed estimate for "
                   "travel and other reimbursements.")
    bp.bullet(doc, "Any discounts offered.")
    bp.bullet(doc, "Pricing for each phase individually and for all three phases as "
                   "a combined engagement.")
    bp.para(doc,
            "Because the three phases share data, models, and project "
            "management, our fee schedule will include an explicit discount "
            "for engaging all three phases as a combined engagement, "
            "presented alongside per-phase pricing.")
    bp.para(doc, [("[Fee schedule to be provided.]", False, True)])


def section_h_v2(doc):
    bp.heading1(doc, "H. Summary", new_page=True)
    bp.para(doc,
            "The Working Group's task is to choose between competing models "
            "with confidence. What Victor-Strategies offers is a record the "
            "Working Group can verify rather than take on faith: more than "
            "25,000 sports bettors interviewed since the repeal of PASPA; "
            "research engagements in more than thirty U.S. states; a "
            "three-year engagement with the Choctaw Nation of Oklahoma that "
            "mirrors this scope; a public Indiana forecast that actual market "
            "results have since confirmed as sound and conservative; the "
            "industry's leading analyst of prediction markets; an economic "
            "modeling platform purpose-built for gaming, backed by more than "
            "165 engagements since 2005; and no relationships with any party "
            "to this decision.")
    bp.para(doc,
            "Behind that record is the team detailed in the next section: "
            "Victor Rocha, who has been educating Indian Country about new "
            "gaming technology for more than 30 years; Gene Johnson, involved "
            "in interactive wagering development since the 1990s; Jay Sarno, "
            "a widely respected analyst with deep casino management and "
            "finance experience; Dustin Gouker on prediction markets; Alfonso "
            "Straffon on sports betting data and revenue modeling; and Kahlil "
            "Philander, among the most widely published economists in the "
            "global gaming industry.")
    bp.para(doc,
            "We appreciate the opportunity to submit this proposal and look forward "
            "to answering any questions you may have. Thank you.")


def section_k_appendix(doc):
    bp.heading1(doc, "K. Appendix — Complete Engagement History", new_page=True)
    bp.para(doc,
            "The following is a chronological history of Victor-Strategies "
            "engagements (and precursor work by its principals) relevant to "
            "interactive wagering, sports betting, and tribal gaming. The "
            "engagements most relevant to this RFP are discussed by phase in "
            "Section A.")
    quals = [
        ("1999–2006. ",
         "Mr. Johnson provided qualitative and quantitative research, analysis, "
         "and consulting services to international iGaming companies operating in "
         "the North American market, covering market sizing, player perceptions, "
         "attributes of attraction, player loyalty and satisfaction, new product "
         "development and testing, and website usability and application "
         "functionality. Results were used to identify opportunities, develop new "
         "games and applications, establish player loyalty and reward programs, "
         "and implement back-end deposit and redemption functions."),
        ("2010. ",
         "While heading EE Johnson Research, Mr. Johnson was contracted by "
         "Spectrum Gaming Group to produce the 2010 Internet Gaming White Paper "
         "for the National Indian Gaming Association — a groundbreaking "
         "examination of the history, trends, and status of global Internet "
         "gaming prior to U.S. legalization that served as an educational "
         "template and operational blueprint for Native American and commercial "
         "casino operators considering domestic iGaming market entry."),
        ("2012. ",
         "Mr. Johnson testified before the United States Senate Committee on "
         "Indian Affairs on the subject of Internet gambling regulation for "
         "Native American tribal governments."),
        ("2012. ",
         "While with Spectrum, Mr. Johnson served as lead consultant to the "
         "Massachusetts State Lottery Commission's Online Products Task Force, "
         "which advised the State Treasurer on the potential for the Lottery — "
         "the most successful in the United States on a per-capita basis — to "
         "offer online wagering products while protecting the interests of its "
         "7,400 retail outlets and the investment value of the soon-to-be-awarded "
         "commercial casino licenses. Mr. Johnson was a principal contributor to "
         "the resulting report, Facing the Lottery's Future: Implications and "
         "Strategies Regarding Internet Sales, which remains a blueprint for how "
         "U.S. state lotteries should approach Internet lottery operations."),
        ("2012. ",
         "Mr. Johnson served as lead consultant and project manager for a major "
         "commercial cruise line exploring onboard iGaming operations in "
         "international waters. Deliverables included an articulated business "
         "model, projected revenues, a technical assessment, and a comprehensive "
         "vendor evaluation, in coordination with multiple international "
         "subcontractors."),
        ("2013. ",
         "Mr. Johnson analyzed the California Internet poker market for a major "
         "Native American tribal casino consortium, including population "
         "statistics, market sizing and demographic profiling of potential "
         "players, projected revenues by poker vertical, optimal taxation rates "
         "based on a survey of legal Internet poker jurisdictions worldwide, and "
         "a global regulatory overview. The results were used to frame policy "
         "decisions by the consortium of tribal governments."),
        ("2014–2016. ",
         "Mr. Johnson designed the content and structure of the opening half-day "
         "iGaming session for the East Coast Gaming Congress and iGaming "
         "Institute in Atlantic City, New Jersey, and moderated each year's panel "
         "on the progress of iGaming in New Jersey, the most successful legal "
         "U.S. jurisdiction."),
        ("2014. ",
         "Mr. Johnson served as lead consultant in a market study and the "
         "development of a market-entry and rollout strategy for a major U.S. "
         "territory exploring iGaming operations."),
        ("2014. ",
         "Mr. Johnson served as Spectrum Gaming Group's lead consultant in two "
         "engagements for casino license applicants under New York's Upstate "
         "Gaming Economic Development Act of 2013, leading a team of subject "
         "matter experts in preparing license submissions for clients in the "
         "Capital and Hudson River districts and participating in the selection "
         "process before the New York State Gaming Commission's Gaming Facility "
         "Location Board."),
        ("2015. ",
         "Mr. Johnson conducted player research for a Native American operator of "
         "five casinos, including two with hotels, in the Great Lakes region. "
         "Objectives included a customer assessment profiling perceptions of the "
         "property ahead of a major refresh, calculating marketing return on "
         "investment, and differentiating the attraction of casino marketing "
         "offers from hotel, food and beverage, and spa offers."),
        ("2016. ",
         "Victor-Strategies conducted a competitive analysis of the Central New "
         "Mexico market for a leading tribal casino operator, measuring the "
         "effect of competitor hotel additions and expansions on visitation to "
         "the client's property in a highly competitive regional marketplace."),
        ("2017. ",
         "Victor-Strategies was engaged by an upstate New York casino operator to "
         "quantify the total available market for sports wagering in New York "
         "State, generate statewide and property-specific revenue estimates, and "
         "develop a mobile and retail sports betting model."),
        ("2018. ",
         "Victor-Strategies authored Sports Betting and Indian Gaming, an "
         "informational white paper on the potential for legalized sports betting "
         "in Indian Country, for a major slot manufacturer. The study was "
         "distributed to select clients shortly after the overturn of PASPA in "
         "May 2018."),
        ("2018. ",
         "Victor-Strategies partnered with Alan Meister of Meister Economic "
         "Consulting to produce the Economic Impact of New Jersey Online Gaming "
         "for iDEA (Interactive Digital Entertainment Association). Updated in "
         "2020 to incorporate six years of regulated-market data, the study "
         "documented New Jersey iGaming's history, operational structure, job "
         "creation, tax revenue, and relationship to land-based casino "
         "performance, and became a widely referenced document in state iGaming "
         "legalization debates."),
        ("2018–2020. ",
         "Victor-Strategies advised a domestic start-up offering micro-betting "
         "products on its U.S. market rollout strategy. The engagement concluded "
         "when the client was acquired by a major U.S. sports betting operator."),
        ("2018–2021. ",
         "Over three years, Victor-Strategies advised the Choctaw Nation of "
         "Oklahoma on the tribe's white-label sports betting initiatives, "
         "delivering a regional market study, an online sports betting model with "
         "a 10-year revenue forecast, a retail sports betting forecast, and an "
         "economic impact analysis. V-S also conducted a survey of more than "
         "3,000 database customers and non-customers across Oklahoma and Texas "
         "profiling betting behavior, attitudes toward legalization, and "
         "demographics — which showed that a significant share of casino players "
         "were actively betting on sports in states where the practice had yet to "
         "be legalized."),
        ("Since 2018. ",
         "Victor-Strategies has partnered with Eilers & Krejcik Gaming on "
         "multiple national and statewide research projects involving online "
         "sports betting market sizing, industry profiles, and gaming consumer "
         "research. Publicly available reports are linked in Section A."),
        ("2019–2022. ",
         "Over three years, Victor-Strategies advised a major Scandinavian sports "
         "betting operator on the development and rollout of its online sports "
         "wagering product in the U.S. market, including competitive intelligence "
         "and assessment of more than a dozen sports betting apps across multiple "
         "states, and consumer research with sports bettors to test the user "
         "interface and develop an Americanized wagering application."),
        ("2019. ",
         "Victor-Strategies conducted a market evaluation of sports betting "
         "feasibility in Arkansas, including statewide and property-specific "
         "handle and revenue forecasts."),
        ("Since 2022. ",
         "Victor-Strategies has advised a tribal gaming client in Minnesota on "
         "the legislative outlook, revenue opportunities, and operational "
         "requirements for legalized sports wagering, developing multiple "
         "market-sizing, handle, and revenue estimates for both mobile and retail "
         "sports betting, adjusted to each session's legislative proposals."),
        ("2023–2025. ",
         "Victor-Strategies executed a two-year engagement with a major "
         "international provider of sports betting data, platforms, and player "
         "account management (PAM) systems, providing strategic advisory services "
         "during U.S. market entry and monthly reports on developments in the "
         "U.S. sports wagering market, with particular attention to tribal "
         "government gaming."),
        ("2024. ",
         "Victor-Strategies collaborated with Alan Meister of Meister Economic "
         "Consulting, Doug Walker, and Dan Waugh of Regulus Partners to produce A "
         "Comprehensive Analysis of NERA's Study on New Jersey's iGaming Economic "
         "Impact — testing two competing economic models' data, assumptions, and "
         "methods line by line, and publishing findings that withstood "
         "adversarial scrutiny from the funded study's sponsors."),
        ("2025. ",
         "Victor-Strategies conducted a sports betting feasibility study for one "
         "of the successful bidders in the New York City casino license "
         "competition, including statewide and property-specific 10-year handle "
         "and revenue forecasts."),
        ("2025. ",
         "Victor-Strategies was engaged by Barclay Damon as an expert witness in "
         "its lawsuit on behalf of the Cayuga Nation against the New York "
         "Lottery, producing an expert witness report on the classification of "
         "lottery offerings as Class III gaming on Indian lands, which is "
         "impermissible under the Indian Gaming Regulatory Act (IGRA)."),
        ("2026. ",
         "Victor-Strategies conducted a sports betting feasibility study for a "
         "tribal gaming client in Minnesota, including statewide and "
         "property-specific five-year handle and revenue forecasts."),
    ]
    for lead, rest in quals:
        bp.bullet_lead(doc, lead, rest)

    refs = doc.add_paragraph()
    bp.set_font(refs.add_run("Reference links: "), size=9.5, bold=True)
    bp.set_font(refs.add_run("2010 Internet Gaming White Paper — "), size=9.5)
    bp.add_hyperlink(refs, "http://www.indiangaming.org/info/alerts/Spectrum-Internet-Paper.pdf")
    bp.set_font(refs.add_run(";  2012 Senate testimony — "), size=9.5)
    bp.add_hyperlink(refs, "http://www.indian.senate.gov/hearing/oversight-hearing-regulation-tribal-gaming-brick-mortar-internet")
    bp.set_font(refs.add_run(";  Massachusetts Lottery report — "), size=9.5)
    bp.add_hyperlink(refs, "http://www.masslottery.com/lib/downloads/leadership/pdfs/SpectrumGamingGroupFinalReport12-4-12Ammended.pdf")
    bp.set_font(refs.add_run(";  New Jersey iGaming study — "), size=9.5)
    bp.add_hyperlink(refs, "https://ideagrowth.org/wp-content/uploads/iDEA-Economic-Impact-of-NJ-iGaming_FULL-REPORT_2019.pdf")
    bp.set_font(refs.add_run(";  NERA rebuttal — "), size=9.5)
    bp.add_hyperlink(refs, "https://ideagrowth.org/wp-content/uploads/iDEA_Comprehensive-Analysis-of-NERA-Study_March-2024.pdf")
    bp.set_font(refs.add_run("."), size=9.5)


# ---------------------------------------------------------------- main

def main():
    doc = Document()
    bp.setup_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top", "bottom"):
        setattr(section, f"{side}_margin", Inches(0.9))
    for side in ("left", "right"):
        setattr(section, f"{side}_margin", Inches(1.0))

    bp.header_footer(doc)
    cover_page_v2(doc)
    bp.table_of_contents(doc)

    executive_summary(doc)
    section_a_v2(doc)
    bp.section_b(doc)
    bp.section_c(doc)
    section_d_v2(doc)
    bp.section_e(doc)
    section_f_v2(doc)
    bp.section_g(doc)
    section_h_v2(doc)
    bp.section_i(doc)
    bp.section_j(doc)
    section_k_appendix(doc)

    props = doc.core_properties
    props.title = ("Online Sports Betting in the State of Wisconsin — "
                   "Data Analysis")
    props.subject = ("Proposal to the Wisconsin Tribal Nations Advisory "
                     "Working Group")
    props.author = "Victor-Strategies"
    props.category = "Proposal"
    props.comments = ("DRAFT 2 — Confidential. Prepared for the Wisconsin "
                      "Tribal Nations Advisory Working Group.")

    bp.enable_update_fields(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
