#!/usr/bin/env python3
"""Build the Victor-Strategies Wisconsin OSB Data Analysis proposal (.docx).

Regenerate after content edits with:
    python3 build_proposal.py        (requires `pip install python-docx`)

Output: VS_Proposal_Online_Sports_Betting_Wisconsin_Data_Analysis_REVISED_061226.docx
"""

import os

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")
OUTPUT = os.path.join(
    HERE, "VS_Proposal_Online_Sports_Betting_Wisconsin_Data_Analysis_REVISED_061226.docx"
)

VS_RED = RGBColor(0xC4, 0x21, 0x27)
VS_DARK = RGBColor(0x3F, 0x3F, 0x3F)
VS_GRAY = RGBColor(0x6E, 0x6E, 0x6E)
BODY_FONT = "Georgia"
HEAD_FONT = "Calibri"


# ---------------------------------------------------------------- low-level helpers

def set_font(run, font=BODY_FONT, size=10.5, bold=False, italic=False, color=None,
             caps=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if caps:
        run.font.all_caps = True
    return run


def add_field(paragraph, instr, dirty=False):
    """Insert a Word field (e.g. PAGE, TOC) as begin/instr/end runs."""
    r1 = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        fld_begin.set(qn("w:dirty"), "true")
    r1._element.append(fld_begin)

    r2 = paragraph.add_run()
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    r2._element.append(instr_el)

    r3 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r3._element.append(fld_end)
    return r2


def add_hyperlink(paragraph, url, text=None, size=9.5):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    rpr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text or url
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade_cell(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def no_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def light_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def fixed_layout(table, col_widths):
    """Force fixed column widths (LibreOffice/Word both respect tblGrid)."""
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    for col, width in zip(table.columns, col_widths):
        col.width = width
    for row in table.rows:
        for cell, width in zip(row.cells, col_widths):
            cell.width = width


def repeat_header_row(table):
    trpr = table.rows[0]._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trpr.append(el)


def keep_row_together(row):
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


def bottom_rule(paragraph, color="C42127", size="12"):
    ppr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    ppr.append(borders)


def enable_update_fields(doc):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


# ---------------------------------------------------------------- style setup

def setup_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = VS_DARK
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15

    h1 = styles["Heading 1"]
    h1.font.name = HEAD_FONT
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = VS_RED
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = HEAD_FONT
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = VS_DARK
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = HEAD_FONT
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.italic = False
    h3.font.color.rgb = VS_GRAY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2"):
        st = styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(10.5)
        st.font.color.rgb = VS_DARK
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.line_spacing = 1.15


# ---------------------------------------------------------------- content helpers

def heading1(doc, text, new_page=False):
    h = doc.add_heading(text, level=1)
    bottom_rule(h, color="D9D9D9", size="8")
    if new_page:
        h.paragraph_format.page_break_before = True
    return h


def para(doc, text_or_runs, style=None, size=10.5, space_after=None,
         align=None):
    """Add a paragraph. text_or_runs is a string or list of
    (text, bold, italic) tuples."""
    p = doc.add_paragraph(style=style)
    if isinstance(text_or_runs, str):
        text_or_runs = [(text_or_runs, False, False)]
    for text, bold, italic in text_or_runs:
        set_font(p.add_run(text), size=size, bold=bold, italic=italic)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def bullet(doc, text_or_runs, level=1):
    style = "List Bullet" if level == 1 else "List Bullet 2"
    return para(doc, text_or_runs, style=style)


def bullet_lead(doc, lead, rest):
    """Bullet with a bold lead-in phrase."""
    return bullet(doc, [(lead, True, False), (rest, False, False)])


def link_bullet(doc, url):
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, url)
    return p


def timeline(doc, text):
    p = doc.add_paragraph()
    set_font(p.add_run("Timeline: "), size=10.5, bold=True)
    set_font(p.add_run(text), size=10.5)
    p.paragraph_format.space_before = Pt(4)
    return p


def superscript(par, text):
    r = par.add_run(text)
    set_font(r, size=10.5)
    r.font.superscript = True
    return r


def bio(doc, name, title_line, photo, paragraphs):
    h = doc.add_heading(level=2)
    set_font(h.add_run(name), font=HEAD_FONT, size=13, bold=True, color=VS_RED,
             caps=True)
    if title_line:
        sub = doc.add_paragraph()
        set_font(sub.add_run(title_line), font=HEAD_FONT, size=10.5, bold=True,
                 italic=True, color=VS_GRAY)
        sub.paragraph_format.space_after = Pt(6)

    if photo:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        no_table_borders(table)
        fixed_layout(table, (Inches(1.75), Inches(4.75)))
        photo_cell, text_cell = table.rows[0].cells
        ph_par = photo_cell.paragraphs[0]
        ph_par.paragraph_format.space_after = Pt(0)
        ph_par.add_run().add_picture(os.path.join(ASSETS, photo), width=Inches(1.45))
        first = True
        for ptext in paragraphs:
            p = text_cell.paragraphs[0] if first else text_cell.add_paragraph()
            first = False
            set_font(p.add_run(ptext), size=10.5)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    else:
        for ptext in paragraphs:
            para(doc, ptext)


# ---------------------------------------------------------------- document parts

def cover_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.add_run().add_picture(os.path.join(ASSETS, "vs-logo.png"), width=Inches(3.6))
    for _ in range(2):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("Online Sports Betting in the State of Wisconsin"),
             font=HEAD_FONT, size=26, bold=True, color=VS_DARK)
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(sub.add_run("Data Analysis"), font=HEAD_FONT, size=18, bold=True,
             color=VS_RED)
    sub.paragraph_format.space_after = Pt(2)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(rule.add_run("—  ❖  —"), font=HEAD_FONT, size=11, color=VS_GRAY)

    doc.add_paragraph()
    prep = doc.add_paragraph()
    prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(prep.add_run("Prepared for"), font=HEAD_FONT, size=12, italic=True,
             color=VS_GRAY)
    prep.paragraph_format.space_after = Pt(2)

    client = doc.add_paragraph()
    client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(client.add_run("Wisconsin Tribal Nations\nAdvisory Working Group"),
             font=HEAD_FONT, size=16, bold=True, color=VS_DARK)

    doc.add_paragraph()
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(date.add_run("June 12, 2026"), font=HEAD_FONT, size=12, color=VS_DARK)

    mark = doc.add_paragraph()
    mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(mark.add_run("DRAFT — CONFIDENTIAL"), font=HEAD_FONT, size=10,
             bold=True, color=VS_RED, caps=True)

    doc.add_page_break()


def table_of_contents(doc):
    h = doc.add_paragraph()
    set_font(h.add_run("Contents"), font=HEAD_FONT, size=16, bold=True,
             color=VS_RED)
    bottom_rule(h, color="D9D9D9", size="8")
    h.paragraph_format.space_after = Pt(10)
    toc_par = doc.add_paragraph()
    add_field(toc_par, r'TOC \o "1-1" \h \z \u', dirty=True)
    note = doc.add_paragraph()
    set_font(note.add_run(
        "If the table of contents does not display, right-click it and choose "
        "“Update Field.”"), size=8.5, italic=True, color=VS_GRAY)
    doc.add_page_break()


def header_footer(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header_par = section.header.paragraphs[0]
    header_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header_par.add_run(
        "Victor-Strategies  |  Wisconsin Online Sports Betting Data Analysis"),
        font=HEAD_FONT, size=8.5, color=VS_GRAY)
    bottom_rule(header_par, color="D9D9D9", size="4")

    footer_par = section.footer.paragraphs[0]
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer_par.add_run("DRAFT — Confidential   |   Page "),
             font=HEAD_FONT, size=8.5, color=VS_GRAY)
    fld_run = add_field(footer_par, "PAGE")
    for r in footer_par.runs[-3:]:
        set_font(r, font=HEAD_FONT, size=8.5, color=VS_GRAY)


# ---------------------------------------------------------------- sections

def section_a(doc):
    heading1(doc, "A. Company Overview and Relevance of Work")

    doc.add_heading("Company Overview", level=2)
    para(doc,
         "Victor-Strategies is a professional advisory firm providing expert "
         "services and critical business insights to the gaming industry, with a "
         "particular focus on Indian Country. We provide governments and industry "
         "leaders with the strategies, tools, data, and expertise essential for "
         "informed decision making and effective operations, supported by "
         "research and analysis capabilities of the highest quality.")
    para(doc,
         "Our business model is built on trusted-advisor relationships through "
         "which clients can draw on our broad network of strategic partners. That "
         "network offers expertise in tribal and commercial gaming — both "
         "land-based and interactive — as well as government policy and media "
         "strategy, and economic impact and development through conventional and "
         "emerging technologies.")
    para(doc,
         "Our principals are recognized thought leaders in the gaming industry "
         "and are committed to delivering maximum value to our clients. Our "
         "founder, Victor Rocha, serves as Conference Chairman of the Indian "
         "Gaming Association while leading Victor-Strategies as its president. He "
         "is also the owner and publisher of Pechanga.net, an online information "
         "resource for Native American tribes across North America, and has been "
         "deeply engaged in the political landscape of U.S. tribal gaming since "
         "1998. Gene Johnson brings more than 35 years of gaming-industry "
         "experience spanning information technology, business process "
         "improvement, quality assurance, and strategic planning and analysis.")
    para(doc,
         "Since its founding in 2016, Victor-Strategies has conducted primary and "
         "secondary research for scores of clients in international markets and "
         "in more than thirty U.S. states. This research includes market sizing, "
         "revenue analysis, feasibility studies, economic impact, player "
         "behavior, gaming consumer research, and new gaming technology.")
    para(doc,
         "Most of our work is confidential and performed on behalf of our "
         "clients. Victor-Strategies has, however, published public-domain "
         "research, including a collaborative study on the economic impact of the "
         "New Jersey iGaming industry, and has conducted a substantial body of "
         "consumer research with casino players and sports betting customers in "
         "both on-premises and interactive wagering environments. Since the "
         "overturn of the Professional and Amateur Sports Protection Act (PASPA) "
         "in May 2018, Victor-Strategies has cumulatively interviewed more than "
         "25,000 sports bettors across its client research projects. Selected "
         "publicly available examples include the following:")
    for url in (
        "https://ideagrowth.org/research/",
        "https://ideagrowth.org/wp-content/uploads/2020/01/Economic-Impact-of-NJ-iGaming_FULL-REPORT_12.19.19.pdf",
        "https://www.in.gov/igc/files/Indiana-SportsBettingReport-Final-Oct18-1.pdf",
        "https://sbcamericas.com/tag/the-more-you-know-the-bettor-defining-us-sports-betting-customer-personas/",
        "http://www.indiangaming.org/info/alerts/Spectrum-Internet-Paper.pdf",
    ):
        link_bullet(doc, url)

    doc.add_heading("Respondent Qualifications", level=2)
    quals = [
        ("1999–2006. ",
         "Mr. Johnson provided qualitative and quantitative research, analysis, "
         "and consulting services to international iGaming companies operating in "
         "the North American market, covering market sizing, player perceptions, "
         "attributes of attraction, player loyalty and satisfaction, new product "
         "development and testing, and website usability and application "
         "functionality. Results were used to identify opportunities, develop new "
         "games and applications, establish player loyalty and reward programs, "
         "and implement back-end deposit and redemption functions."),
        ("2010. ",
         "While heading EE Johnson Research, Mr. Johnson was contracted by "
         "Spectrum Gaming Group to produce the 2010 Internet Gaming White Paper "
         "for the National Indian Gaming Association — a groundbreaking "
         "examination of the history, trends, and status of global Internet "
         "gaming prior to U.S. legalization that served as an educational "
         "template and operational blueprint for Native American and commercial "
         "casino operators considering domestic iGaming market entry."),
        ("2012. ",
         "Mr. Johnson testified before the United States Senate Committee on "
         "Indian Affairs on the subject of Internet gambling regulation for "
         "Native American tribal governments."),
        ("2012. ",
         "While with Spectrum, Mr. Johnson served as lead consultant to the "
         "Massachusetts State Lottery Commission's Online Products Task Force, "
         "which advised the State Treasurer on the potential for the Lottery — "
         "the most successful in the United States on a per-capita basis — to "
         "offer online wagering products while protecting the interests of its "
         "7,400 retail outlets and the investment value of the soon-to-be-awarded "
         "commercial casino licenses. Mr. Johnson was a principal contributor to "
         "the resulting report, Facing the Lottery's Future: Implications and "
         "Strategies Regarding Internet Sales, which remains a blueprint for how "
         "U.S. state lotteries should approach Internet lottery operations."),
        ("2012. ",
         "Mr. Johnson served as lead consultant and project manager for a major "
         "commercial cruise line exploring onboard iGaming operations in "
         "international waters. Deliverables included an articulated business "
         "model, projected revenues, a technical assessment, and a comprehensive "
         "vendor evaluation, in coordination with multiple international "
         "subcontractors."),
        ("2013. ",
         "Mr. Johnson analyzed the California Internet poker market for a major "
         "Native American tribal casino consortium, including population "
         "statistics, market sizing and demographic profiling of potential "
         "players, projected revenues by poker vertical, optimal taxation rates "
         "based on a survey of legal Internet poker jurisdictions worldwide, and "
         "a global regulatory overview. The results were used to frame policy "
         "decisions by the consortium of tribal governments."),
        ("2014–2016. ",
         "Mr. Johnson designed the content and structure of the opening half-day "
         "iGaming session for the East Coast Gaming Congress and iGaming "
         "Institute in Atlantic City, New Jersey, and moderated each year's panel "
         "on the progress of iGaming in New Jersey, the most successful legal "
         "U.S. jurisdiction."),
        ("2014. ",
         "Mr. Johnson served as lead consultant in a market study and the "
         "development of a market-entry and rollout strategy for a major U.S. "
         "territory exploring iGaming operations."),
        ("2014. ",
         "Mr. Johnson served as Spectrum Gaming Group's lead consultant in two "
         "engagements for casino license applicants under New York's Upstate "
         "Gaming Economic Development Act of 2013, leading a team of subject "
         "matter experts in preparing license submissions for clients in the "
         "Capital and Hudson River districts and participating in the selection "
         "process before the New York State Gaming Commission's Gaming Facility "
         "Location Board."),
        ("2015. ",
         "Mr. Johnson conducted player research for a Native American operator of "
         "five casinos, including two with hotels, in the Great Lakes region. "
         "Objectives included a customer assessment profiling perceptions of the "
         "property ahead of a major refresh, calculating marketing return on "
         "investment, and differentiating the attraction of casino marketing "
         "offers from hotel, food and beverage, and spa offers."),
        ("2016. ",
         "Victor-Strategies conducted a competitive analysis of the Central New "
         "Mexico market for a leading tribal casino operator, measuring the "
         "effect of competitor hotel additions and expansions on visitation to "
         "the client's property in a highly competitive regional marketplace."),
        ("2017. ",
         "Victor-Strategies was engaged by an upstate New York casino operator to "
         "quantify the total available market for sports wagering in New York "
         "State, generate statewide and property-specific revenue estimates, and "
         "develop a mobile and retail sports betting model."),
        ("2018. ",
         "Victor-Strategies authored Sports Betting and Indian Gaming, an "
         "informational white paper on the potential for legalized sports betting "
         "in Indian Country, for a major slot manufacturer. The study was "
         "distributed to select clients shortly after the overturn of PASPA in "
         "May 2018."),
        ("2018. ",
         "Victor-Strategies partnered with Alan Meister of Meister Economic "
         "Consulting to produce the Economic Impact of New Jersey Online Gaming "
         "for iDEA (Interactive Digital Entertainment Association). Updated in "
         "2020 to incorporate six years of regulated-market data, the study "
         "documented New Jersey iGaming's history, operational structure, job "
         "creation, tax revenue, and relationship to land-based casino "
         "performance, and became a widely referenced document in state iGaming "
         "legalization debates."),
        ("2018–2020. ",
         "Victor-Strategies advised a domestic start-up offering micro-betting "
         "products on its U.S. market rollout strategy. The engagement concluded "
         "when the client was acquired by a major U.S. sports betting operator."),
        ("2018–2021. ",
         "Over three years, Victor-Strategies advised the Choctaw Nation of "
         "Oklahoma on the tribe's white-label sports betting initiatives, "
         "delivering a regional market study, an online sports betting model with "
         "a 10-year revenue forecast, a retail sports betting forecast, and an "
         "economic impact analysis. V-S also conducted a survey of more than "
         "3,000 database customers and non-customers across Oklahoma and Texas "
         "profiling betting behavior, attitudes toward legalization, and "
         "demographics — which showed that a significant share of casino players "
         "were actively betting on sports in states where the practice had yet to "
         "be legalized."),
        ("Since 2018. ",
         "Victor-Strategies has partnered with Eilers & Krejcik Gaming on "
         "multiple national and statewide research projects involving online "
         "sports betting market sizing, industry profiles, and gaming consumer "
         "research. Publicly available reports are linked at the beginning of "
         "this section."),
        ("2019–2022. ",
         "Over three years, Victor-Strategies advised a major Scandinavian sports "
         "betting operator on the development and rollout of its online sports "
         "wagering product in the U.S. market, including competitive intelligence "
         "and assessment of more than a dozen sports betting apps across multiple "
         "states, and consumer research with sports bettors to test the user "
         "interface and develop an Americanized wagering application."),
        ("2019. ",
         "Victor-Strategies conducted a market evaluation of sports betting "
         "feasibility in Arkansas, including statewide and property-specific "
         "handle and revenue forecasts."),
        ("Since 2022. ",
         "Victor-Strategies has advised a tribal gaming client in Minnesota on "
         "the legislative outlook, revenue opportunities, and operational "
         "requirements for legalized sports wagering, developing multiple "
         "market-sizing, handle, and revenue estimates for both mobile and retail "
         "sports betting, adjusted to each session's legislative proposals."),
        ("2023–2025. ",
         "Victor-Strategies executed a two-year engagement with a major "
         "international provider of sports betting data, platforms, and player "
         "account management (PAM) systems, providing strategic advisory services "
         "during U.S. market entry and monthly reports on developments in the "
         "U.S. sports wagering market, with particular attention to tribal "
         "government gaming."),
        ("2024. ",
         "Victor-Strategies collaborated with Alan Meister of Meister Economic "
         "Consulting, Doug Walker, and Dan Waugh of Regulus Partners to produce A "
         "Comprehensive Analysis of NERA's Study on New Jersey's iGaming Economic "
         "Impact, a rebuttal to a flawed report funded by the Campaign for Fairer "
         "Gambling that disputed the findings of our earlier New Jersey economic "
         "impact study."),
        ("2025. ",
         "Victor-Strategies conducted a sports betting feasibility study for one "
         "of the successful bidders in the New York City casino license "
         "competition, including statewide and property-specific 10-year handle "
         "and revenue forecasts."),
        ("2025. ",
         "Victor-Strategies was engaged by Barclay Damon as an expert witness in "
         "its lawsuit on behalf of the Cayuga Nation against the New York "
         "Lottery, producing an expert witness report on the classification of "
         "lottery offerings as Class III gaming on Indian lands, which is "
         "impermissible under the Indian Gaming Regulatory Act (IGRA)."),
        ("2026. ",
         "Victor-Strategies conducted a sports betting feasibility study for a "
         "tribal gaming client in Minnesota, including statewide and "
         "property-specific five-year handle and revenue forecasts."),
    ]
    for lead, rest in quals:
        bullet_lead(doc, lead, rest)

    refs = doc.add_paragraph()
    set_font(refs.add_run("Reference links: "), size=9.5, bold=True)
    set_font(refs.add_run(
        "2010 Internet Gaming White Paper — "), size=9.5)
    add_hyperlink(refs, "http://www.indiangaming.org/info/alerts/Spectrum-Internet-Paper.pdf")
    set_font(refs.add_run(";  2012 Senate testimony — "), size=9.5)
    add_hyperlink(refs, "http://www.indian.senate.gov/hearing/oversight-hearing-regulation-tribal-gaming-brick-mortar-internet")
    set_font(refs.add_run(";  Massachusetts Lottery report — "), size=9.5)
    add_hyperlink(refs, "http://www.masslottery.com/lib/downloads/leadership/pdfs/SpectrumGamingGroupFinalReport12-4-12Ammended.pdf")
    set_font(refs.add_run(";  New Jersey iGaming study — "), size=9.5)
    add_hyperlink(refs, "https://ideagrowth.org/wp-content/uploads/iDEA-Economic-Impact-of-NJ-iGaming_FULL-REPORT_2019.pdf")
    set_font(refs.add_run(";  NERA rebuttal — "), size=9.5)
    add_hyperlink(refs, "https://ideagrowth.org/wp-content/uploads/iDEA_Comprehensive-Analysis-of-NERA-Study_March-2024.pdf")
    set_font(refs.add_run("."), size=9.5)

    para(doc,
         "Victor-Strategies and the associates assembled for this proposal have "
         "demonstrated extensive experience in sports betting market analysis. We "
         "are familiar with IGRA Class III compact structures and tribal "
         "revenue-sharing frameworks. We have no current relationships with any "
         "of the 11 Wisconsin tribes and no current financial relationship with "
         "proponents of either model under review.")


def section_b(doc):
    heading1(doc, "B. Phase I — Market Assessment and Data Validation", new_page=True)
    para(doc,
         "Victor-Strategies and its team of expert associates will develop an "
         "independent statewide online/mobile sports betting model forecasting "
         "handle and revenue over a 10-year period. The revenue estimate will "
         "incorporate assumptions about expected future trends and the "
         "cannibalization effects of new gaming verticals, including social "
         "sports sites and — most importantly — prediction markets.")

    doc.add_heading("Approach and Methodology", level=2)
    doc.add_heading("Wisconsin Mobile Sports Betting Market Size (RFP §8.1.1)",
                    level=3)
    para(doc,
         "We will use historical data from states where sports betting is legal "
         "to develop comparison metrics for Wisconsin, qualifying those "
         "comparisons with demographic statistics including adult population, "
         "median and average household income, and consumer purchasing power. The "
         "most important comparison metric will be revenue per adult (RPA) — "
         "gross gaming revenue (GGR) per adult (21+) — which can be compared "
         "directly with states having similar legalization regimes, populations, "
         "and income demographics.")
    para(doc,
         "Alfonso Straffon, an expert associate on this engagement, maintains a "
         "proprietary database of online sports betting statistics used by Wall "
         "Street investment firms. This data will support our independent revenue "
         "forecast and serve as a check on the existing revenue models under "
         "review.")
    para(doc,
         "Demand-based modeling will be used to estimate potential retail sports "
         "betting revenue based on current and future population demographics, "
         "gambling penetration, propensity to gamble, and the expected effects of "
         "increased regional competition for each of the 11 Indian tribes of "
         "Wisconsin.")
    para(doc,
         "Specifically, Victor-Strategies will develop reasonable assumptions for "
         "the likely evolution of the statewide sports wagering market over the "
         "next ten years and chart its likely trajectory. Because the Wisconsin "
         "legalization model is a version of hub-and-spoke featuring tribal "
         "exclusivity, we will apply a discount when comparing to commercial "
         "online sports betting (OSB) states where national brands face "
         "unrestricted competition.")
    para(doc,
         "In accordance with the RFP instructions, our independent model will "
         "include a sensitivity analysis at multiple discount levels to establish "
         "a range of plausible market outcomes. We will also review the potential "
         "tribal-operator-only market size with no promotional restrictions "
         "versus a 30 percent promotional restriction.")

    doc.add_heading("Data Accuracy Review (RFP §8.1.2)", level=3)
    para(doc,
         "As instructed in the RFP, Victor-Strategies will conduct a data "
         "accuracy review, evaluating the accuracy, completeness, and reliability "
         "of all submitted raw data and identifying inconsistencies, unsupported "
         "assumptions, and material gaps. Our independent revenue model and "
         "proprietary database will provide the basis for evaluating the data and "
         "assumptions underlying the other models.")
    para(doc,
         "Ideally, the Phase I independent statewide revenue model will be "
         "completed before we review the existing models' accuracy and "
         "assumptions; client timelines will determine our actual sequencing.")

    doc.add_heading("Phase I Deliverables (RFP §8.1.3)", level=2)
    bullet(doc, "Market-sizing summary with GGR estimates and documented "
                "assumptions.")
    bullet(doc, "Written assessment of data accuracy, gaps, and limitations in "
                "all submitted models.")
    bullet(doc, "Adjusted Wisconsin GGR forecast with sensitivity analysis.")
    timeline(doc, "Completed within 45 days of contract signing.")


def section_c(doc):
    heading1(doc, "C. Phase II — Revenue Share Model Review and Comparison")

    doc.add_heading("Approach and Methodology", level=2)
    doc.add_heading("Comparative Analysis (RFP §8.2.1)", level=3)
    para(doc,
         "Victor-Strategies will develop parallel comparisons of all reviewed "
         "models to determine the most equitable and beneficial revenue-sharing "
         "arrangement. These comparisons will cover revenue definitions (GGR "
         "versus net gaming revenue), expense allocations, risk assumptions, "
         "scalability, and market sensitivity. We will analyze variations in "
         "contribution structures — including alternative percentages and "
         "breakpoints such as those featured in the Red Cliff framework — to "
         "assess their effects on distribution levels, operator risk, and "
         "operational costs. We will also identify the financial risks and costs "
         "borne by operating tribes under each model to which non-operating "
         "tribes are not exposed.")

    doc.add_heading("All-Model Financial Analysis", level=3)
    para(doc, "As instructed in the RFP, the Victor-Strategies team will:")
    bullet(doc,
           "Verify that the net gaming revenue (NGR) definition — GGR less "
           "federal excise tax and promotional credits, with no further "
           "deductions — is industry-standard and protective against operator "
           "circumvention.")
    bullet(doc,
           "Quantify per-tribe capital exposure under each model: upfront capital "
           "expenditures and annual operating expenses until breakeven, and the "
           "percentage of profit retained for reinvestment versus distributed.")
    bullet(doc,
           "Evaluate the potential market impacts of multiple apps operating "
           "under one operator.")
    bullet(doc,
           "Provide a scenario analysis of the three potential brand structures: "
           "(1) a newly created tribal brand; (2) an existing Wisconsin tribal "
           "gaming brand used as the communal brand; or (3) a nationally "
           "recognized non-Wisconsin gaming brand licensed to the joint venture. "
           "For each scenario we will assess startup costs, ongoing fees, the "
           "impact on per-tribe distributions, and the effect on tribal "
           "sovereignty and brand control, and compare each scenario financially "
           "to the B2B model.")

    doc.add_heading("GGR Estimation and 10-Year Comparison (RFP §8.2.2)", level=3)
    para(doc, "As instructed in the RFP, Victor-Strategies will:")
    bullet(doc,
           "Estimate total GGR for the Wisconsin market based on Phase I "
           "validated data and market assumptions, clearly identifying all "
           "assumptions used.")
    bullet(doc,
           "Present a side-by-side, 10-year, per-tribe economic comparison of all "
           "models using consistent GGR assumptions from Phase I.")

    doc.add_heading("Phase II Deliverables (RFP §8.2.3)", level=2)
    bullet(doc, "Revenue model comparison matrix covering all submitted models.")
    bullet(doc, "Written assessment of strengths, weaknesses, and risk allocation "
                "by model.")
    bullet(doc, "Any additional model recommendations that Victor-Strategies "
                "would propose for Wisconsin.")
    timeline(doc, "Completed within 45 days of contract signing.")


def section_d(doc):
    heading1(doc, "D. Phase II Option — Statewide Economic Impact Study", new_page=True)

    doc.add_heading("Background", level=2)
    para(doc,
         "Economic impact analysis was not specifically requested in the "
         "Wisconsin Tribal Nations request for proposal. During the question "
         "period, however, Victor-Strategies asked whether statewide economic "
         "impact was an element of any of the models under review, and whether "
         "the consultant should include a statewide economic impact component in "
         "its independent revenue model. The answer was affirmative, so we "
         "include an economic impact component in this response. It is offered as "
         "an additional option, priced separately from the other phases of our "
         "proposal.")

    doc.add_heading("Approach and Methodology", level=2)
    para(doc,
         "If included, this macroeconomic analysis will estimate the total "
         "statewide economic impact of online sports wagering as forecast by the "
         "recommended sports betting revenue model. The impact of this new "
         "economic vertical will be measured across several dimensions: economic "
         "output; value added at the state level, or gross state product (GSP); "
         "full-time-equivalent (FTE) employment; employee compensation; and "
         "non-gaming tax revenue. Results will be reported as direct, indirect, "
         "and induced impacts to the State of Wisconsin and, by inference, to the "
         "tribes.")
    para(doc,
         "Economic impact is a measure of the spending and employment associated "
         "with a business, a sector of the economy, a specific project (such as "
         "the construction of a new facility), or a change in government policy "
         "or regulation. It can be measured in several ways, most commonly the "
         "dollar value of output produced and the person-years (full-time "
         "equivalents) of employment generated. These figures assess the gross "
         "level of activity or expenditure; they are not net measures that weigh "
         "benefits against costs, but they are useful in developing an "
         "appreciation of businesses, projects, investments, and economic "
         "sectors. In our modeling we do account for the shifts in consumption "
         "that are clearest to define — namely, shifts by Wisconsin residents "
         "from other household goods and services to sports wagering.")

    doc.add_heading("The GEMS Modeling Platform", level=3)
    para(doc,
         "Our calculations will be performed with GEMS (Gaming Economic Modeling "
         "System), a proprietary input-output (IO) modeling platform developed by "
         "GP Consulting and purpose-built for the gaming industry. GEMS covers "
         "all 50 states and is constructed from federal economic accounts: state "
         "IO multipliers derived from U.S. EPA state input-output models, "
         "employment and wage coefficients from the Bureau of Labor Statistics "
         "Quarterly Census of Employment and Wages, and industry detail from the "
         "Bureau of Economic Analysis (BEA).")
    para(doc,
         "GEMS addresses a structural weakness of standard, off-the-shelf "
         "economic impact software for gaming applications. Conventional "
         "platforms model casinos within a blended industry sector that combines "
         "gambling with golf courses, fitness centers, and amusement parks — "
         "activities with fundamentally different labor and cost structures — "
         "which systematically misstates the employment and value-added effects "
         "of gaming operations. GEMS isolates gambling-specific multipliers from "
         "the BEA detail accounts and adds dedicated online-gambling coefficients "
         "that conventional models lack entirely. For an online sports betting "
         "analysis, this produces materially more accurate estimates than "
         "generic, blended-sector multipliers.")
    para(doc,
         "GEMS produces transparent, replicable estimates of output, value added, "
         "employment, wages, and state and local fiscal effects, decomposed into "
         "direct, indirect, and induced impacts using Type I and Type II "
         "(Leontief) multipliers. It incorporates direct employment data where "
         "available and enables rapid, consistent comparison of alternative "
         "revenue scenarios — a useful feature when evaluating the alternative "
         "market structures contemplated in this engagement.")

    doc.add_heading("Measures of Impact", level=3)
    para(doc,
         "The statewide economic impact of sports wagering expansion will be "
         "measured across the following dimensions:")
    bullet(doc, "Economic output")
    bullet(doc, "Value added, or gross state product (GSP)")
    bullet(doc, "Full-time-equivalent (FTE) employment")
    bullet(doc, "Employee compensation")
    bullet(doc, "Non-gaming tax revenue (production, payroll, and household "
                "taxes)")
    para(doc,
         "We estimate economic activity through three layers of effects — "
         "direct, indirect, and induced:")
    bullet_lead(doc, "Direct economic impact ",
                "is the employment and economic output attributable to the "
                "operation and management of the sportsbook within the operator's "
                "business. This includes the jobs located directly at the "
                "operator, together with the professional staff, management, "
                "marketing, and other workers required to run the platform.")
    bullet_lead(doc, "Indirect economic impact ",
                "is the employment, value added, and economic output created in "
                "industries that supply goods and services to the operator — for "
                "example, a local IT company that installs hardware, or an "
                "in-state office supplier that furnishes back-of-house spaces.")
    bullet_lead(doc, "Induced economic impact ",
                "is the employment, value added, and economic output generated by "
                "the spending of individuals employed directly or indirectly by "
                "the operator — for example, operator employees consuming "
                "services such as a meal at a Wisconsin restaurant.")
    bullet_lead(doc, "Total economic impact ",
                "is the sum of direct, indirect, and induced effects. The "
                "multiplier (indirect and induced) impacts represent the maximum "
                "potential stimulus to the economy from gaming-related business "
                "activity.")

    doc.add_heading("Catalytic Impacts", level=3)
    p = para(doc,
             "Total economic impacts are conventionally defined as the sum of "
             "direct, indirect, and induced effects. Some sectors, however, are "
             "also capable of producing catalytic impacts — economic growth "
             "enabled by another sector, industry, or firm. Aviation is often "
             "cited as a source of catalytic impacts because it allows trade that "
             "would otherwise not occur, such as air cargo transport of "
             "perishable goods.")
    superscript(p, "1")
    para(doc,
         "In Wisconsin, the expansion of sports betting could similarly "
         "facilitate catalytic impacts. If sports betting attracts new "
         "entrepreneurs to the state, for example, the industry may spur the "
         "development of a new set of businesses that would not otherwise have "
         "considered expanding there.")
    para(doc,
         "Under a typical, conservative economic impact methodology, such "
         "potential catalytic impacts are not measured. We considered including "
         "catalytic estimates, but they would require significant assumptions "
         "about innovation and consumer behavior that would materially increase "
         "the margin of error on our projections. We are nonetheless comfortable "
         "asserting that adopting a more business-friendly set of policies will "
         "typically have a net positive effect on surrounding businesses, and we "
         "note that our figures and approach are conservative: there is reason to "
         "believe the Wisconsin economy would experience impacts beyond those "
         "specifically measured in this study.")

    doc.add_heading("Evidence of Impacts on Existing Gaming Businesses", level=3)
    p = para(doc,
             "A common public policy question is whether new gaming revenue will "
             "cannibalize existing businesses in the area — in particular, other "
             "gaming businesses. The question is especially important when the "
             "potentially cannibalized business produces higher margins or "
             "public revenue. This topic has received extensive study in the "
             "academic literature, which has consistently pointed to a "
             "non-negative effect: offline gaming businesses are estimated to "
             "realize positive revenue gains from increased sports betting "
             "participation when the systems can be horizontally integrated.")
    superscript(p, "2")
    para(doc,
         "An important test of whether economic impact figures can be trusted "
         "for policymaking is whether a negative impact is occurring elsewhere "
         "in the economy. As noted above, the gaming economics literature "
         "suggests that aggregate effects on existing gaming businesses will be "
         "positive or nil (not cannibalistic). While we expect this to be the "
         "case for total impacts, we make explicit assumptions about changes in "
         "household discretionary spending by Wisconsin residents: we assume "
         "that revenue from Wisconsin residents is diverted from other spending, "
         "based on existing spending patterns and scenario-based analyses.")

    doc.add_heading("Economic Impact Deliverables", level=2)
    bullet(doc,
           "A separate section of our management summary report detailing the "
           "statewide direct, indirect, induced, and total economic impacts of "
           "an online sports betting industry across the tribal and commercial "
           "economies of Wisconsin.")
    timeline(doc, "Priced separately as an optional addition to Phase II; "
                  "schedule to be aligned with the Phase II deliverables.")

    notes = doc.add_paragraph()
    set_font(notes.add_run("Notes"), font=HEAD_FONT, size=10, bold=True,
             color=VS_GRAY)
    notes.paragraph_format.space_before = Pt(14)
    notes.paragraph_format.keep_with_next = True
    bottom_rule(notes, color="D9D9D9", size="4")
    n1 = doc.add_paragraph()
    n1.paragraph_format.keep_together = True
    n1.paragraph_format.keep_with_next = True
    superscript(n1, "1 ")
    set_font(n1.add_run(
        "Tretheway, M. (2010). Economic Impacts of Aviation: Catalytic Impacts. "
        "ACED Conference. "), size=9)
    add_hyperlink(n1, "http://www.intervistas.com/downloads/presentations/Economic_Impacts_of_Aviation_Catalytic_Impacts_MTR_20Sep2010.pdf", size=9)
    n1.paragraph_format.space_after = Pt(3)
    n2 = doc.add_paragraph()
    n2.paragraph_format.keep_together = True
    superscript(n2, "2 ")
    set_font(n2.add_run(
        "For a more thorough discussion of this literature, see Philander, K., "
        "Abarbanel, B., & Repetti, T. (2015). Consumer spending in the gaming "
        "industry: Evidence of complementary demand in casino and online venues. "
        "International Gambling Studies. "), size=9)
    add_hyperlink(n2, "https://www.tandfonline.com/doi/abs/10.1080/14459795.2015.1042002", size=9)


def section_e(doc):
    heading1(doc, "E. Phase III — Prediction Markets", new_page=True)

    doc.add_heading("Approach and Methodology", level=2)
    doc.add_heading("Prediction Market Analysis (RFP §8.3.1)", level=3)
    para(doc, "As instructed in the RFP, the Victor-Strategies team will:")
    bullet(doc,
           "Analyze how prediction markets are currently affecting the states in "
           "which they operate — specifically their effect on licensed sports "
           "betting handle, GGR, and tribal revenue.")
    bullet(doc,
           "Assess the impact of prediction market growth on the Wisconsin market "
           "projections from Phase I, including any required adjustments to GGR "
           "estimates.")
    bullet(doc,
           "Provide a reasoned opinion on the trajectory of prediction markets "
           "based on current data and regulatory developments, and on what that "
           "trajectory means for the long-term viability of Wisconsin tribal "
           "sports betting exclusivity.")

    doc.add_heading("Phase III Deliverables (RFP §8.3.2)", level=2)
    bullet(doc, "Written analysis of prediction market impacts on comparable "
                "markets.")
    bullet(doc, "Assessment of prediction market risk to Wisconsin tribal sports "
                "betting projections.")
    timeline(doc, "Completed within 60 days of contract signing.")


def section_f(doc):
    heading1(doc, "F. Financial Proposal (RFP §8.4)")
    para(doc,
         "Per the RFP, respondents must provide a detailed fee schedule "
         "including:")
    bullet(doc, "Estimated hours by staff level, with hourly rates and total "
                "cost.")
    bullet(doc, "Out-of-pocket expenses, including a not-to-exceed estimate for "
                "travel and other reimbursements.")
    bullet(doc, "Any discounts offered.")
    bullet(doc, "Pricing for each phase individually and for all three phases as "
                "a combined engagement.")
    para(doc, [("[Fee schedule to be provided.]", False, True)])


def section_g(doc):
    heading1(doc, "G. Key Questions the Report Must Answer (RFP §9)")
    para(doc,
         "The RFP requires the final report to provide a documented, "
         "evidence-based response to each of the following questions, stating "
         "explicitly where information is missing from either model. The table "
         "below restates each question and identifies where our proposed scope "
         "addresses it.")

    rows = [
        ("Is the Wisconsin market size projection appropriate — conservative, "
         "reasonable, or aggressive — relative to comparable states?",
         "Phase I (Section B): independent market model benchmarked on revenue "
         "per adult in comparable states, with sensitivity analysis."),
        ("Is the tribal exclusivity discount supported by comparable markets, "
         "and is the phase-down schedule consistent with observed market "
         "maturation?",
         "Phase I (Section B): hub-and-spoke exclusivity discount benchmarked "
         "against commercial OSB states, tested at multiple discount levels."),
        ("How does the number of operators in the market affect total handle "
         "and GGR per adult in comparable U.S. sports betting markets?",
         "Phases I and II (Sections B and C): comparable-state RPA analysis and "
         "evaluation of multiple apps under one operator."),
        ("What are the financial implications of each potential brand "
         "structure?",
         "Phase II (Section C): three-scenario brand analysis compared "
         "financially to the B2B model."),
        ("How are prediction markets currently impacting licensed sports "
         "betting revenue in comparable states, and what does that mean for "
         "Wisconsin tribal projections?",
         "Phase III (Section E): prediction market analysis and risk assessment."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    light_table_borders(table)
    repeat_header_row(table)
    hdr = table.rows[0].cells
    for i, label in enumerate(("RFP Question", "Where Addressed")):
        shade_cell(hdr[i], "C42127")
        p = hdr[i].paragraphs[0]
        set_font(p.add_run(label), font=HEAD_FONT, size=10, bold=True,
                 color=RGBColor(0xFF, 0xFF, 0xFF))
        p.paragraph_format.space_after = Pt(2)
    for q, a in rows:
        row = table.add_row()
        keep_row_together(row)
        for cell, text in zip(row.cells, (q, a)):
            p = cell.paragraphs[0]
            set_font(p.add_run(text), size=9.5)
            p.paragraph_format.space_after = Pt(2)
    fixed_layout(table, (Inches(3.6), Inches(3.4)))


def section_h(doc):
    heading1(doc, "H. Summary", new_page=True)
    para(doc,
         "Victor-Strategies brings a strong set of relevant skills and "
         "experience to this research requirement, and we have assembled a "
         "remarkable team of experts to address it. Victor Rocha has been "
         "educating Indian Country about new gaming technology for more than 30 "
         "years. Gene Johnson has been involved in interactive wagering "
         "development since the 1990s. Jay Sarno is a widely respected data "
         "analyst with a rich background in casino management, finance, and data "
         "analysis. Dustin Gouker is arguably the industry's most informed "
         "expert on prediction markets. Alfonso Straffon has a wealth of "
         "experience in the online sports betting industry and is highly skilled "
         "in statistics and the creation of revenue models. Kahlil Philander is "
         "among the most experienced economists and most widely published "
         "academics in the global gaming industry.")
    para(doc,
         "We appreciate the opportunity to submit this proposal and look forward "
         "to answering any questions you may have. Thank you.")


def section_i(doc):
    heading1(doc, "I. Our Team")

    bio(doc, "Gene Johnson",
        "Executive Vice President, Victor-Strategies", None, [
        "Gene Johnson has more than 30 years of unique experience in the gaming "
        "industry, bringing together technology, market research, strategic "
        "planning and analysis, and casino marketing — all framed through the "
        "lens of the customer. He began working in Atlantic City casinos in 1989 "
        "and in 1997 founded EE Johnson Research, a marketing research and "
        "consulting firm specializing in the casino industry with a particular "
        "focus on gambling motivation and behavior. In 2012 he became a senior "
        "executive at an internationally recognized casino gaming consultancy. "
        "In 2016 he founded Gaming Knowledge Partners to serve casino gaming "
        "clients through an expert network of strategic partners with deep "
        "knowledge of gambling behavior, operations, marketing, and property "
        "management; later that year he co-founded Victor-Strategies with "
        "partners Rob Miller and Victor Rocha.",
        "With Victor-Strategies, Mr. Johnson extends this expert network "
        "globally to provide clients with the highest quality advisory services "
        "aimed at improving operations, increasing revenue, and building "
        "sustainable economic growth.",
        "Over his career, Mr. Johnson has managed and conducted numerous "
        "qualitative and quantitative studies with gamblers of all types, as "
        "well as extensive consumer research for a broad array of corporate, "
        "government, and non-profit clients. He has testified before the U.S. "
        "Senate, the Federal Trade Commission, and multiple state and tribal "
        "governments; has written articles featured in many industry "
        "publications; and speaks regularly at major gaming industry conferences "
        "on a wide range of subjects.",
        "Mr. Johnson holds a BA from Washington College in Maryland and an MBA "
        "from the University of Phoenix. He has earned Certified Quality Analyst "
        "certification from the Quality Assurance Institute and is a member of "
        "the American Society for Quality, the American Marketing Association, "
        "and the Qualitative Research Consultants Association.",
        ])

    bio(doc, "Victor Rocha",
        "President, Victor-Strategies", "rocha.png", [
        "Victor Rocha is the president, chief motivator, and inspiration of "
        "Victor-Strategies. A respected voice in Indian Country and a frequent "
        "contributor to national media publications, Mr. Rocha is a successful "
        "Native American Internet entrepreneur with a long-term commitment to "
        "bringing information and opportunity to Indian Country. His conferences "
        "have received high praise for the quality of their content and "
        "speakers, and for bringing political and gaming leaders together with "
        "the latest trends and technology. He is a recognized leader in emerging "
        "gaming technologies who keeps Indian gaming informed about advances in "
        "technology, interactive wagering, and new forms of gambling.",
        "Mr. Rocha is also the owner and editor of Pechanga.net, the premier "
        "source of news and information on Indian gaming. Involved in the "
        "politics of Indian gaming since 1998, he has earned numerous awards for "
        "his work, including the National Indian Gaming Association's 2002 "
        "Outstanding Contribution to Indian Country, VCAT's 2001 Catalyst Award, "
        "Global Gaming Business Magazine's “40 Under 40,” Raving's 2012 Casino "
        "Marketing Lifetime Achievement Award, and the AGA's Lifetime "
        "Achievement Award for Gaming Marketing. He currently serves as "
        "Conference Chair for the National Indian Gaming Association's annual "
        "Indian Gaming Tradeshow and Mid-Year Conference.",
        ])

    bio(doc, "Jay Sarno",
        "Strategic Partner", "sarno.jpg", [
        "Jay Sarno is a Victor-Strategies strategic partner widely known in the "
        "gaming industry for his analytical skills — building rigorous "
        "quantitative models, streamlining processes, and providing solutions "
        "that answer a wide range of client needs.",
        "A twenty-five-year business veteran, Jay has broad practical expertise "
        "across operational and applied areas including traditional corporate "
        "finance, off-Wall Street capital management and investment analysis, "
        "sales management, compliance, business intelligence software "
        "development, and statistical analysis. He has conducted feasibility "
        "studies in the Midwest and market analyses of the Iowa regional gaming "
        "market.",
        "Jay has held positions ranging from Vice President of Business "
        "Development and Corporate Compliance Officer to Assistant Casino "
        "General Manager and Financial Officer, with experience at Advanced "
        "Casino Systems, Hollywood Casino Corporation, the Sands Casino in "
        "Atlantic City, and the Golden Nugget Casino/Resort in Las Vegas.",
        "Jay received his BS in Finance from the University of Nevada, Las "
        "Vegas and his MBA from Sacred Heart University in Connecticut. He has "
        "taught undergraduate- and graduate-level business coursework in the "
        "School of Professional Studies at Stockton University in New Jersey.",
        ])

    bio(doc, "Dustin Gouker",
        "Strategic Partner", "gouker.jpg", [
        "Dustin Gouker has been in the gambling industry in one way or another "
        "for two decades, beginning with playing poker and creating poker "
        "content. He spent much of his early career as a newspaper writer and "
        "editor before joining a startup in 2015 to lead the Legal Sports "
        "Report blog in its early days covering the daily fantasy sports "
        "industry. From there he helped build LSR into the go-to source of "
        "information on the sports betting industry — along with a network of "
        "other gambling content sites that became leaders in the space. The "
        "startup was acquired in 2017 by Catena Media, where Dustin became head "
        "of content and eventually a vice president for North America.",
        "He left that role at the start of 2023 and now offers consulting "
        "services through Closing Line Consulting and publishes two "
        "newsletters, The Closing Line and The Event Horizon, reaching nearly "
        "1,500 subscribers. Dustin has been cited over the years by numerous "
        "media outlets, including The Washington Post, ESPN, Yahoo Finance, and "
        "Axios.",
        ])

    bio(doc, "Alfonso Straffon",
        "Strategic Partner", None, [
        "Alfonso brings a unique mix of relevant experience to the project. He "
        "began his career in the sports betting industry in his native Costa "
        "Rica, where he worked in risk management and trading for eight years. "
        "In 2008 he left the industry to pursue a career in investment "
        "management, primarily in roles as a bond trader and research analyst. "
        "In 2018, following the repeal of PASPA, he began offering his services "
        "as an independent consultant and is regarded as a highly reliable "
        "source of data and industry analytics.",
        "He holds a Bachelor's degree in Business Administration from the "
        "University of Costa Rica and an MBA from Wake Forest University, and "
        "is currently pursuing a Master's in Statistics at NC State University. "
        "He speaks fluent English and Spanish.",
        ])

    bio(doc, "Kahlil Philander, PhD",
        "Lead Economist — GP Consulting", "philander.jpg", [
        "Dr. Kahlil Philander is an economist and gaming researcher "
        "specializing in the analysis of large-scale economic development "
        "initiatives, regulatory policy, and consumer behavior in the gaming "
        "industry. He is a tenured Associate Professor at Washington State "
        "University's Carson College of Business and holds an honorary "
        "appointment in the School of Psychology at the University of Sydney. "
        "With nearly 20 years of applied research experience across academia, "
        "industry, and government, he offers a distinctive blend of technical "
        "expertise and policy insight, particularly in contexts involving "
        "tourism, entertainment, and community impact.",
        "His research spans regional economic forecasting, taxation policy, and "
        "the socioeconomic outcomes of the gaming industry, and has been funded "
        "by government agencies and private-sector clients. Across more than "
        "165 consulting engagements since 2005, Dr. Philander has led economic "
        "impact, market, and fiscal analyses of gaming for state, tribal, and "
        "national governments and for major operators. He previously served as "
        "Director of Social Responsibility at the British Columbia Lottery "
        "Corporation, where he oversaw the GameSense program, and as Director "
        "of Research at the University of Nevada, Las Vegas International "
        "Gaming Institute. In those roles he collaborated with stakeholders "
        "ranging from legislators and municipal planners to health officials "
        "and corporate executives to evaluate the economic, social, and "
        "regulatory implications of major capital projects.",
        "Dr. Philander holds a PhD in Hospitality Administration from the "
        "University of Nevada, Las Vegas, with a dissertation on the economic "
        "impact of casino tax policy; an MA in Economics from the University of "
        "Toronto; and a BCom (Honors) in Finance and Economics from the "
        "University of British Columbia. His portfolio includes more than 40 "
        "peer-reviewed publications in top-tier journals such as the Journal of "
        "Policy Modeling, the Journal of Gambling Studies, and Tourism "
        "Management, alongside numerous industry reports. He has provided "
        "expert testimony to governmental bodies and consulted for major gaming "
        "companies and governments on economic modeling and policy strategy; "
        "his commentary has been featured by CNBC, the Financial Times, and "
        "Wired.",
        ])

    doc.add_heading("Notable Works", level=3)
    works = [
        "Philander, K., Abarbanel, B., & Repetti, T. (2015). Consumer spending "
        "in the gaming industry: Evidence of complementary demand in casino and "
        "online venues. International Gambling Studies.",
        "Philander, K., Wimmer, B., Bernhard, B., & Singh, A. U.S. casino "
        "revenue taxes and short-run labor outcomes. Journal of Policy "
        "Modeling.",
        "Philander, K. (2014). Specific or ad valorem? A theory of casino "
        "taxation. Tourism Economics.",
        "Philander, K. (2017). Entry fees as a responsible gambling tool: An "
        "economic analysis. UNLV Gaming Research & Review Journal.",
        "Philander, K. (2019). Regional impacts of casino availability on "
        "gambling problems: Evidence from the Canadian Community Health Survey. "
        "Tourism Management.",
        "Philander, K., Tabri, N., Wood, R., & Wohl, M. (2022). Retail access "
        "and addictive consumption: Evidence from casino gambling. "
        "International Gambling Studies.",
        "Philander, K., & Roe, S. (2013). The impact of wage rate growth on "
        "tourism competitiveness. Tourism Economics.",
    ]
    for w in works:
        bullet(doc, w)


def section_j(doc):
    heading1(doc, "J. About GP Consulting", new_page=True)
    para(doc,
         "GP Consulting provides economic analysis and policy research on "
         "gambling markets for governments, tribal nations, regulators, and "
         "commercial operators. The firm's work spans statewide market and "
         "economic impact studies, gaming tax and fiscal policy analysis, "
         "revenue and demand forecasting, responsible gambling program "
         "development and audit, and consulting and testifying expert support "
         "in U.S. and international proceedings. Engagements range from "
         "single-jurisdiction impact assessments to multi-year advisory "
         "relationships with national governments and newly established "
         "regulatory authorities.")
    para(doc,
         "At the core of the firm's quantitative work is the Gaming Economic "
         "Modeling System (GEMS), a proprietary input-output platform covering "
         "all 50 states and built on federal economic accounts (BEA, BLS) with "
         "land-based and online gaming-specific multipliers. GEMS is "
         "purpose-built for gaming in a way standard economic impact software "
         "is not. Off-the-shelf platforms model casinos inside a blended "
         "industry sector that combines gambling with golf courses, fitness "
         "centers, and amusement parks — activities with fundamentally "
         "different labor and cost structures — which systematically misstates "
         "the employment and value-added effects of a gaming property. GEMS "
         "isolates gambling-specific multipliers from federal detail accounts "
         "and adds dedicated online-gambling coefficients that conventional "
         "models lack entirely, producing materially more accurate results for "
         "both land-based and online operations.")
    para(doc,
         "GEMS produces transparent, replicable estimates of output, value "
         "added, employment, wages, and state and local fiscal effects, "
         "decomposed into direct, indirect, and induced impacts. It "
         "incorporates direct employment data where available and enables "
         "rapid, consistent comparison of alternative revenue scenarios.")
    para(doc,
         "GP Consulting has completed more than 165 engagements since 2005, "
         "including economic impact and market studies for state, tribal, and "
         "national governments and for major gaming operators. Beyond economic "
         "analysis, the firm brings deep regulatory and policy expertise.")

    doc.add_heading("Selected Engagements", level=2)
    engagements = [
        ("Economic Impact Analysis — Choctaw Nation of Oklahoma. ",
         "Lead Economist for analysis of the Nation's gaming operations, with "
         "Victor-Strategies."),
        ("Economic Impact Analysis and Market Forecast — Washington State "
         "Gambling Commission. ",
         "Principal investigator for gaming market analysis and revenue "
         "forecasting."),
        ("Economic Analysis and Forecasting of Sports Betting — Seminole Hard "
         "Rock. ",
         "Analysis supporting Florida sports betting strategy."),
        ("Economic Impact Study of Sports Betting — Indiana Gaming Commission. ",
         "Lead Economist for the statewide economic and fiscal impact analysis "
         "of legal sports wagering, including alternative tax scenarios."),
        ("Economic and Market Study of iGaming — West Virginia Lottery. ",
         "Lead Economist for market and economic analysis of internet gaming "
         "legalization."),
        ("Economic Study of the National Gaming and Lottery Sector — Government "
         "of The Bahamas. ",
         "Lead Economist for an engagement spanning gaming tax reform."),
        ("Economic Impact Modeling — Wynn Resorts. ",
         "Lead Economist for development impact modeling."),
        ("Economic and Social Impact Study of Japanese Integrated Resorts — "
         "US-Japan Business Council, via UNLV International Gaming Institute. ",
         ""),
        ("Economic Impact Study of Online and Retail Gaming — MGM Resorts "
         "International. ",
         "Published economic impact analysis of online and land-based "
         "operations."),
        ("Economic Research and Regulatory Advisory Services — General "
         "Commercial Gaming Regulatory Authority, United Arab Emirates. ",
         "Research Advisor to a newly established federal gaming regulator."),
    ]
    for lead, rest in engagements:
        if rest:
            bullet_lead(doc, lead, rest)
        else:
            bullet(doc, [(lead.rstrip(), True, False)])


# ---------------------------------------------------------------- main

def main():
    doc = Document()
    setup_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top", "bottom"):
        setattr(section, f"{side}_margin", Inches(0.9))
    for side in ("left", "right"):
        setattr(section, f"{side}_margin", Inches(1.0))

    header_footer(doc)
    cover_page(doc)
    table_of_contents(doc)

    section_a(doc)
    section_b(doc)
    section_c(doc)
    section_d(doc)
    section_e(doc)
    section_f(doc)
    section_g(doc)
    section_h(doc)
    section_i(doc)
    section_j(doc)

    props = doc.core_properties
    props.title = ("Online Sports Betting in the State of Wisconsin — "
                   "Data Analysis")
    props.subject = ("Proposal to the Wisconsin Tribal Nations Advisory "
                     "Working Group")
    props.author = "Victor-Strategies"
    props.category = "Proposal"
    props.comments = ("DRAFT — Confidential. Prepared for the Wisconsin Tribal "
                      "Nations Advisory Working Group.")

    enable_update_fields(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
